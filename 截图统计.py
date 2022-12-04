# %%
# 当使用腾讯文档收集截图时，输出格式良好的报告文件

# %%
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from math import ceil
import pandas as pd
import requests
from PIL import Image

# %%
# 当前文件目录
path = "D:\\python"
# 主标题
headtext = "腾讯文档截图统计"
# 打开的EXCEL，格式为{学号,姓名,截图URL}
fileopen = path + "\\20221204.xls"
# fileopen="D:\\python\\20221204.xls"
sheetname = "Sheet1"
# 保存文件位置
filesave = path + "\\" + headtext + ".docx"
# filesave = "C:\\Users\\admin\\Desktop\\腾讯文档截图统计.docx"
# 保留的截图位置
pictemp = path + "\\pictemp.jpg"
# pictemp = "C:\\Users\\admin\\Desktop\\pictemp.jpg"
# 截图过大压缩（1为不压缩，范围[0,1]）
zipratio = 0.28


# %%
df = pd.read_excel(fileopen, sheet_name=sheetname,
                   index_col=0, dtype=str).values
total = len(df)

# %%
headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/96.0.4664.45 Safari/537.36'
}

# %%
document = Document()
# 添加标题
head = document.add_paragraph()
run = head.add_run(headtext)
# 修改字体
head.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 字体居中
run.font.size = Pt(20)          # 字体大小
run.bold = True                 # 字体是否加粗
run.font.name = 'Times New Roman'           # 控制是西文时的字体
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')  # 控制是中文时的字体

document.add_paragraph()

# %%
table = document.add_table(rows=ceil(total/5), cols=5)
table.style = 'TableGrid'

# %%

for j in range(0, total):
    # 姓名
    run = table.cell(j//5, j % 5).paragraphs[0].add_run(df[j][1]+"\n")
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)

    # 学号
    run = table.cell(j//5, j % 5).paragraphs[0].add_run(df[j][0]+"\n")
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)

    # 获取截图图片
    r = requests.get(df[j][2], headers=headers)
    f = open(pictemp, 'wb')
    f.write(r.content)
    f.close()

    # 压缩图片
    im = Image.open(pictemp)
    w, h = im.size
    im_ss = im.resize((int(w*zipratio), int(h*zipratio)))
    im_ss = im.convert('RGB')
    im_ss.save(pictemp)

    # 导入截图图片
    run = table.cell(j//5, j % 5).paragraphs[0].add_run()
    picture = run.add_picture(pictemp)
    picture.height = Cm(4)
    picture.width = Cm(2)

# %%
document.save(filesave)
